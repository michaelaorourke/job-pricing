"""
Document processing service for PDF, DOCX, and text files
"""

import io
import PyPDF2
from docx import Document
from typing import Optional
from fastapi import UploadFile
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process various document formats"""

    async def extract_text(self, file: UploadFile) -> Optional[str]:
        """Extract text from uploaded file"""

        filename = file.filename.lower()
        content = await file.read()

        try:
            if filename.endswith('.pdf'):
                return self._extract_pdf(content)
            elif filename.endswith(('.docx', '.doc')):
                return self._extract_docx(content)
            elif filename.endswith('.txt'):
                return content.decode('utf-8', errors='ignore')
            else:
                # Try to decode as text
                return content.decode('utf-8', errors='ignore')

        except Exception as e:
            logger.error(f"Document processing error: {e}")
            return None

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"

            return text.strip()

        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return ""

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)

            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)

            return "\n".join(text)

        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ""

    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""

        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = [line.strip() for line in lines if line.strip()]
        text = '\n'.join(cleaned_lines)

        # Remove special characters but keep essential punctuation
        import re
        text = re.sub(r'[^\w\s\-.,;:!()\[\]{}\'\"\/]', ' ', text)
        text = re.sub(r'\s+', ' ', text)

        return text.strip()